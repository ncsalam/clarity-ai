import os
from flask import Blueprint, request, jsonify
from werkzeug.utils import secure_filename
import pypdf
import docx
import json
from datetime import datetime
from pydantic import ValidationError

from .main import db
from .models import (
    Document, 
    Requirement, 
    Tag, 
    ProjectSummary, 
    UserProfile,
    AmbiguityAnalysis, 
    AmbiguousTerm, 
    ClarificationHistory,
    ContradictionAnalysis
)
from .rag_service import (
    process_and_store_document,
    generate_project_requirements,
    generate_project_summary,
    delete_document_from_rag,
    _save_summary_to_db
)
from .auth_service import get_roles_permissions_config, require_auth
from .ambiguity_service import AmbiguityService
from .schemas import (
    AmbiguityAnalyzeRequest,
    AmbiguityAnalyzeRequirementRequest,
    AmbiguityBatchAnalyzeRequest,
    ClarificationSubmitRequest,
    ReportExportRequest,
    LexiconAddRequest,
    ContradictionAnalysisSchema
)
from .validation_utils import rate_limiter
from .contradiction_analysis_service import ContradictionAnalysisService 
from .edge_case_service import EdgeCaseService

api_bp = Blueprint('api', __name__, url_prefix='/api')
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'md', 'json'}


@api_bp.route('/')
def index():
    return jsonify({"message": "Welcome to the Clarity AI API!"})


@api_bp.route('/health')
def health_check():
    """Health check endpoint for monitoring and CI/CD"""
    return jsonify({
        "status": "healthy",
        "service": "Clarity AI API",
        "timestamp": datetime.utcnow().isoformat()
    }), 200

# --- Health Check Endpoints ---


@api_bp.route('/health/supertokens', methods=['GET'])
def supertokens_health_check():
    """
    Health check endpoint to verify SuperTokens connectivity and configuration.
    """
    try:
        # Test SuperTokens connectivity by making a simple request to SuperTokens Core
        import requests
        from flask import current_app

        # Get SuperTokens connection URI from config
        supertokens_uri = current_app.config.get(
            'connection_uri', 'http://localhost:3567')

        try:
            # Test connectivity to SuperTokens Core
            response = requests.get(f"{supertokens_uri}/hello", timeout=5)

            if response.status_code == 200:
                return jsonify({
                    "status": "healthy",
                    "service": "SuperTokens",
                    "message": "SuperTokens Core is connected and responding",
                    "core_url": supertokens_uri,
                    "timestamp": datetime.utcnow().isoformat()
                })
            else:
                return jsonify({
                    "status": "unhealthy",
                    "service": "SuperTokens",
                    "message": f"SuperTokens Core responded with status {response.status_code}",
                    "core_url": supertokens_uri,
                    "timestamp": datetime.utcnow().isoformat()
                }), 503

        except requests.exceptions.RequestException as e:
            return jsonify({
                "status": "unhealthy",
                "service": "SuperTokens",
                "message": f"SuperTokens Core connection failed: {str(e)}",
                "core_url": supertokens_uri,
                "timestamp": datetime.utcnow().isoformat()
            }), 503

    except Exception as e:
        return jsonify({
            "status": "error",
            "service": "SuperTokens",
            "message": f"Health check failed: {str(e)}",
            "timestamp": datetime.utcnow().isoformat()
        }), 500


@api_bp.route('/health/database', methods=['GET'])
def database_health_check():
    """
    Health check endpoint to verify database connectivity.
    """
    try:
        # Test database connectivity by running a simple query
        from .main import db
        result = db.session.execute(db.text('SELECT 1')).fetchone()

        if result:
            return jsonify({
                "status": "healthy",
                "service": "Database",
                "message": "Database connection is working",
                "timestamp": datetime.utcnow().isoformat()
            })
        else:
            return jsonify({
                "status": "unhealthy",
                "service": "Database",
                "message": "Database query returned no result",
                "timestamp": datetime.utcnow().isoformat()
            }), 503

    except Exception as e:
        return jsonify({
            "status": "unhealthy",
            "service": "Database",
            "message": f"Database connectivity error: {str(e)}",
            "timestamp": datetime.utcnow().isoformat()
        }), 503


@api_bp.route('/health/full', methods=['GET'])
def full_health_check():
    """
    Comprehensive health check endpoint that verifies all system components.
    """
    from flask import current_app

    health_status = {
        "overall_status": "healthy",
        "timestamp": datetime.utcnow().isoformat(),
        "services": {}
    }

    # Check API health
    health_status["services"]["api"] = {
        "status": "healthy",
        "message": "API is running"
    }

    # Check database health
    try:
        from .main import db
        result = db.session.execute(db.text('SELECT 1')).fetchone()
        if result:
            health_status["services"]["database"] = {
                "status": "healthy",
                "message": "Database connection is working"
            }
        else:
            health_status["services"]["database"] = {
                "status": "unhealthy",
                "message": "Database query returned no result"
            }
            health_status["overall_status"] = "degraded"
    except Exception as e:
        health_status["services"]["database"] = {
            "status": "unhealthy",
            "message": f"Database connectivity error: {str(e)}"
        }
        health_status["overall_status"] = "unhealthy"

    # Check SuperTokens health
    try:
        import requests

        # Get SuperTokens connection URI from config
        supertokens_uri = current_app.config.get(
            'connection_uri', 'http://localhost:3567')

        # Test connectivity to SuperTokens Core
        response = requests.get(f"{supertokens_uri}/hello", timeout=5)

        if response.status_code == 200:
            health_status["services"]["supertokens"] = {
                "status": "healthy",
                "message": "SuperTokens Core is connected and responding",
                "core_url": supertokens_uri,
                "roles_initialized": current_app.config.get('SUPERTOKENS_INITIALIZED', False),
                "config_valid": current_app.config.get('SUPERTOKENS_CONFIG_VALID', False)
            }
        else:
            health_status["services"]["supertokens"] = {
                "status": "unhealthy",
                "message": f"SuperTokens Core responded with status {response.status_code}",
                "core_url": supertokens_uri,
                "roles_initialized": current_app.config.get('SUPERTOKENS_INITIALIZED', False),
                "config_valid": current_app.config.get('SUPERTOKENS_CONFIG_VALID', False)
            }
            health_status["overall_status"] = "degraded"

    except requests.exceptions.RequestException as e:
        health_status["services"]["supertokens"] = {
            "status": "unhealthy",
            "message": f"SuperTokens Core connection failed: {str(e)}",
            "roles_initialized": current_app.config.get('SUPERTOKENS_INITIALIZED', False),
            "config_valid": current_app.config.get('SUPERTOKENS_CONFIG_VALID', False)
        }
        health_status["overall_status"] = "degraded"
    except Exception as e:
        health_status["services"]["supertokens"] = {
            "status": "unhealthy",
            "message": f"SuperTokens health check failed: {str(e)}",
            "roles_initialized": current_app.config.get('SUPERTOKENS_INITIALIZED', False),
            "config_valid": current_app.config.get('SUPERTOKENS_CONFIG_VALID', False)
        }
        health_status["overall_status"] = "degraded"

    # Determine overall status code
    if health_status["overall_status"] == "healthy":
        return jsonify(health_status)
    elif health_status["overall_status"] == "degraded":
        return jsonify(health_status), 200
    else:
        return jsonify(health_status), 503


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def parse_file_content(file_storage):
    filename = file_storage.filename
    extension = filename.rsplit('.', 1)[1].lower()
    content = ""
    file_storage.seek(0)

    if extension in ['txt', 'md']:
        content = file_storage.read().decode('utf-8')
    elif extension == 'json':
        try:
            json_data = json.load(file_storage)
            content = json.dumps(json_data, indent=2)
        except json.JSONDecodeError:
            raise ValueError("Invalid JSON file.")
    elif extension == 'pdf':
        pdf_reader = pypdf.PdfReader(file_storage)
        for page in pdf_reader.pages:
            content += page.extract_text()
    elif extension == 'docx':
        doc = docx.Document(file_storage)
        for para in doc.paragraphs:
            content += para.text + '\n'

    return content

# --- NEW: Get all documents ---


@api_bp.route('/documents', methods=['GET'])
@require_auth(["documents:read"])
def get_documents():
    """
    Fetches all documents from the database, filtered by owner_id if user is authenticated.
    """
    try:
        # Get current user ID from authenticated session
        from flask import g
        current_user_id = g.user_id

        # Filter documents by owner_id for authenticated user
        documents = Document.query.filter_by(
            owner_id=current_user_id).order_by(Document.created_at.desc()).all()

        results = [
            {
                "id": doc.id,
                "filename": doc.filename,
                "created_at": doc.created_at.isoformat() if doc.created_at else None
            }
            for doc in documents
        ]
        return jsonify(results)
    except Exception as e:
        print(f"An error occurred while fetching documents: {str(e)}")
        return jsonify({"error": "Failed to fetch documents"}), 500

@api_bp.route('/upload', methods=['POST'])
@require_auth(["documents:write"])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"error": "No file part in the request"}), 400

    file = request.files['file']

    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        new_document = None

        try:
            content = parse_file_content(file)

            # Get current user ID from authenticated session
            from flask import g
            current_user_id = g.user_id

            new_document = Document(
                filename=filename, content=content, owner_id=current_user_id)
            db.session.add(new_document)
            db.session.commit()

            process_and_store_document(new_document)
            # Return the new document object, matching the /documents GET route
            return jsonify({
                "message": "File uploaded and processed successfully",
                "document": {
                    "id": new_document.id,
                    "filename": new_document.filename,
                    "created_at": new_document.created_at.isoformat() if new_document.created_at else None
                }
            }), 201
        
        except Exception as e:
            print(f"An error occurred during file processing: {str(e)}")
            db.session.rollback()
            return jsonify({"error": f"Failed to process file: {str(e)}"}), 500
        
    return jsonify({"error": "File type not allowed"}), 400

# --- NEW: Delete a document ---


@api_bp.route('/documents/<int:document_id>', methods=['DELETE'])
@require_auth(["documents:write"])
def delete_document(document_id):
    """
    Deletes a document from the database and the RAG vector store, with user access validation.
    """
    try:
        # Get current user ID from authenticated session
        from flask import g
        current_user_id = g.user_id

        # Find document with user access validation
        document = Document.query.filter_by(
            id=document_id, owner_id=current_user_id).first()

        if not document:
            return jsonify({"error": "Document not found or access denied"}), 404

        # 1. Delete from RAG first.
        delete_document_from_rag(document_id)

        # 2. Delete the document object from the session.
        db.session.delete(document)

        # 3. Commit the transaction
        db.session.commit()

        return jsonify({"message": f"Document ID {document_id} and associated data deleted."}), 200

    except Exception as e:
        db.session.rollback()
        print(f"An error occurred during document deletion: {str(e)}")
        return jsonify({"error": f"Failed to delete document: {str(e)}"}), 500


@api_bp.route('/requirements/generate', methods=['POST'])
@require_auth(["requirements:write"])
def trigger_requirements_generation():
    """
    Triggers a full regeneration of requirements from user's documents if authenticated.
    This clears existing user requirements first.
    """
    try:
        # Get current user ID from authenticated session
        from flask import g
        current_user_id = g.user_id

        total_generated = generate_project_requirements(
            owner_id=current_user_id)
        return jsonify({
            "message": f"Successfully generated {total_generated} new requirements."
        })
    except Exception as e:
        print(f"An error occurred during requirements generation: {str(e)}")
        return jsonify({"error": f"Failed to generate requirements: {str(e)}"}), 500

@api_bp.route('/requirements', methods=['GET'])
@require_auth(["requirements:read"])
def get_requirements():
    """
    Fetches all requirements from the database, filtered by owner_id if user is authenticated,
    including their associated tags and the filename of their source document.
    Optimized with eager loading to prevent N+1 query problems.
    """
    try:
        # Get current user ID from authenticated session
        from flask import g
        current_user_id = g.user_id

        # Use optimized query with eager loading
        from .database_optimization import get_requirements_with_relations
        requirements = get_requirements_with_relations(current_user_id)

        results = []
        for req in requirements:
            results.append({
                "id": req.id,
                "req_id": req.req_id,
                "title": req.title,
                "description": req.description,
                "status": req.status,
                "priority": req.priority,
                "requirement_type": req.requirement_type,
                "stakeholders": req.stakeholders,
                "source_document_filename": req.source_document.filename if req.source_document else None,
                "tags": [{"id": tag.id, "name": tag.name} for tag in req.tags]
            })

        return jsonify(results)

    except Exception as e:
        print(f"An error occurred while fetching requirements: {str(e)}")
        return jsonify({"error": "Failed to fetch requirements"}), 500

# --- NEW: Get requirements count ---


@api_bp.route('/requirements/count', methods=['GET'])
@require_auth(["requirements:read"])
def get_requirements_count():
    """
    Fetches a simple count of requirements, filtered by owner_id if user is authenticated.
    """
    try:
        # Get current user ID from authenticated session
        from flask import g
        current_user_id = g.user_id

        # Filter count by owner_id for authenticated user
        count = db.session.query(Requirement.id).filter_by(
            owner_id=current_user_id).count()

        return jsonify({"count": count})
    except Exception as e:
        print(f"An error occurred while fetching requirements count: {str(e)}")
        return jsonify({"error": "Failed to fetch requirements count"}), 500

@api_bp.route('/summary', methods=['GET'])
@require_auth(["summary:read"])
def get_summary():
    """
    Fetches the latest generated project summary (as a JSON string) 
    from the database.
    """
    try:
        from flask import g
        current_user_id = g.user_id

        latest_summary = ProjectSummary.query.filter_by(
            owner_id=current_user_id).order_by(ProjectSummary.created_at.desc()).first()

        if latest_summary:
            # --- MODIFIED: Send the raw JSON string from the DB ---
            # The frontend will be responsible for parsing this.
            return jsonify({
                "summary": latest_summary.content, 
                "created_at": latest_summary.created_at.isoformat() if latest_summary.created_at else None
            })
        else:
            return jsonify({
                "summary": None, # Send null instead of a string
                "created_at": None
            }), 404

    except Exception as e:
        print(f"An error occurred while fetching summary: {str(e)}")
        return jsonify({"error": "Failed to fetch summary"}), 500

@api_bp.route('/summary/generate', methods=['POST'])
@require_auth(["summary:write"])
def trigger_summary_generation():
    """
    Triggers synchronous generation of a new summary.
    Returns the new summary as a JSON string.
    """
    try:
        from flask import g
        current_user_id = g.user_id

        # 1. Call the generation function, returns a Pydantic object
        summary_object = generate_project_summary(owner_id=current_user_id)
        
        # 2. Convert the Pydantic object to a JSON string
        summary_json_string = summary_object.model_dump_json()

        # 3. Save the new JSON string to the database
        _save_summary_to_db(summary_json_string, current_user_id)

        # 4. Return the new JSON string
        return jsonify({
            "message": "Summary generated successfully.",
            "summary": summary_json_string,
            "created_at": datetime.utcnow().isoformat()
        }), 201

    except Exception as e:
        db.session.rollback()
        print(f"An error occurred during summarization: {str(e)}")
        return jsonify({"error": f"Failed to summarize project: {str(e)}"}), 500


# --- Profile Management Endpoints ---

def assign_default_role_to_user(user_id, email):
    """
    Assigns a default role to a new user based on business logic.
    For now, assigns 'pilot-user' role to all new users.
    """
    try:
        # Import here to avoid circular imports
        from supertokens_python.recipe import userroles
        import asyncio

        # Determine role based on business logic
        # For now, all new users get 'pilot-user' role
        default_role = 'pilot-user'

        # Create event loop if needed
        try:
            loop = asyncio.get_event_loop()
        except RuntimeError:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)

        # Assign role to user
        async def assign_role():
            try:
                result = await userroles.add_role_to_user(
                    tenant_id="public",
                    user_id=user_id,
                    role=default_role
                )
                return result
            except Exception as e:
                print(f"Error assigning role to user {user_id}: {str(e)}")
                return None

        result = loop.run_until_complete(assign_role())

        if result:
            print(
                f"Successfully assigned role '{default_role}' to user {user_id}")
            return default_role
        else:
            print(f"Failed to assign role to user {user_id}")
            return None

    except Exception as e:
        print(f"Error in role assignment for user {user_id}: {str(e)}")
        return None


@api_bp.route('/auth/profile', methods=['POST'])
def create_profile():
    """
    Creates a new user profile.
    Expects JSON payload with user profile data.
    """
    try:
        data = request.get_json()

        # Validate required fields
        required_fields = ['user_id', 'email', 'first_name',
                           'last_name', 'company', 'job_title']
        for field in required_fields:
            if not data.get(field):
                return jsonify({"error": f"Missing required field: {field}"}), 400

        # Check if profile already exists
        existing_profile = UserProfile.query.filter_by(
            user_id=data['user_id']).first()
        if existing_profile:
            return jsonify({"error": "Profile already exists for this user"}), 409

        # Create new profile
        new_profile = UserProfile(
            user_id=data['user_id'],
            email=data['email'],
            first_name=data['first_name'],
            last_name=data['last_name'],
            company=data['company'],
            job_title=data['job_title'],
            remaining_tokens=data.get(
                'remaining_tokens', 5)  # Default to 5 tokens
        )

        db.session.add(new_profile)
        db.session.commit()

        # Assign default role to the new user
        assigned_role = assign_default_role_to_user(
            data['user_id'], data['email'])

        return jsonify({
            "message": "Profile created successfully",
            "profile": {
                "id": new_profile.id,
                "user_id": new_profile.user_id,
                "email": new_profile.email,
                "first_name": new_profile.first_name,
                "last_name": new_profile.last_name,
                "company": new_profile.company,
                "job_title": new_profile.job_title,
                "remaining_tokens": new_profile.remaining_tokens,
                "created_at": new_profile.created_at.isoformat() if new_profile.created_at else None,
                "assigned_role": assigned_role
            }
        }), 201

    except Exception as e:
        db.session.rollback()
        print(f"An error occurred during profile creation: {str(e)}")
        return jsonify({"error": f"Failed to create profile: {str(e)}"}), 500


@api_bp.route('/auth/profile', methods=['GET'])
def get_profile():
    """
    Retrieves user profile by user_id from query parameters.
    """
    try:
        user_id = request.args.get('user_id')
        if not user_id:
            return jsonify({"error": "user_id parameter is required"}), 400

        profile = UserProfile.query.filter_by(user_id=user_id).first()
        if not profile:
            return jsonify({"error": "Profile not found"}), 404

        return jsonify({
            "profile": {
                "id": profile.id,
                "user_id": profile.user_id,
                "email": profile.email,
                "first_name": profile.first_name,
                "last_name": profile.last_name,
                "company": profile.company,
                "job_title": profile.job_title,
                "remaining_tokens": profile.remaining_tokens,
                "created_at": profile.created_at.isoformat() if profile.created_at else None,
                "updated_at": profile.updated_at.isoformat() if profile.updated_at else None
            }
        })

    except Exception as e:
        print(f"An error occurred while fetching profile: {str(e)}")
        return jsonify({"error": f"Failed to fetch profile: {str(e)}"}), 500


@api_bp.route('/profile', methods=['GET'])
@require_auth(["profile:read"])
def get_current_user_profile():
    """
    Retrieves the current authenticated user's profile.
    """
    try:
        # Get current user ID from authenticated session
        from flask import g
        current_user_id = g.user_id

        profile = UserProfile.query.filter_by(user_id=current_user_id).first()
        if not profile:
            return jsonify({"error": "Profile not found"}), 404

        # Return profile in the format expected by frontend
        return jsonify({
            "user_id": profile.user_id,
            "email": profile.email,
            "metadata": {
                "first_name": profile.first_name,
                "last_name": profile.last_name,
                "user_profile": {
                    "company": profile.company,
                    "job_title": profile.job_title
                },
                "user_token": {
                    "remaining_tokens": profile.remaining_tokens
                }
            }
        })

    except Exception as e:
        print(f"An error occurred while fetching profile: {str(e)}")
        return jsonify({"error": f"Failed to fetch profile: {str(e)}"}), 500


@api_bp.route('/profile', methods=['PUT'])
@require_auth(["profile:write"])
def update_current_user_profile():
    """
    Updates the current authenticated user's profile.
    """
    try:
        # Get current user ID from authenticated session
        from flask import g
        current_user_id = g.user_id

        data = request.get_json()
        if not data:
            return jsonify({"error": "No data provided"}), 400

        # Extract metadata from request
        metadata = data.get('metadata', {})
        user_profile = metadata.get('user_profile', {})

        # Validate required fields
        required_fields = {
            'first_name': metadata.get('first_name'),
            'last_name': metadata.get('last_name'),
            'company': user_profile.get('company'),
            'job_title': user_profile.get('job_title')
        }

        for field, value in required_fields.items():
            if not value or str(value).strip() == '':
                return jsonify({"error": f"Missing required field: {field}"}), 400

        # Find existing profile
        profile = UserProfile.query.filter_by(user_id=current_user_id).first()
        if not profile:
            return jsonify({"error": "Profile not found"}), 404

        # Update profile fields
        profile.first_name = required_fields['first_name']
        profile.last_name = required_fields['last_name']
        profile.company = required_fields['company']
        profile.job_title = required_fields['job_title']

        # Update remaining tokens if provided
        user_token = metadata.get('user_token', {})
        if 'remaining_tokens' in user_token:
            profile.remaining_tokens = user_token['remaining_tokens']

        db.session.commit()

        # Return updated profile in the format expected by frontend
        return jsonify({
            "user_id": profile.user_id,
            "email": profile.email,
            "metadata": {
                "first_name": profile.first_name,
                "last_name": profile.last_name,
                "user_profile": {
                    "company": profile.company,
                    "job_title": profile.job_title
                },
                "user_token": {
                    "remaining_tokens": profile.remaining_tokens
                }
            }
        })

    except Exception as e:
        db.session.rollback()
        print(f"An error occurred while updating profile: {str(e)}")
        return jsonify({"error": f"Failed to update profile: {str(e)}"}), 500

# --- Edge Case Generation Endpoint ---  # NEW

@api_bp.route('/requirements/<int:requirement_id>/edge-cases', methods=['POST'])  # NEW
@require_auth(["requirements:write"])  # NEW
def generate_edge_cases(requirement_id):  # NEW
    """
    Generate edge test cases for a specific requirement.  # NEW
    Returns a simple JSON payload with a list of edge case descriptions.  # NEW
    """  # NEW
    try:  # NEW
        from flask import g  # NEW
        current_user_id = g.user_id  # NEW

        # Optional: reuse the same rate limiter style as ambiguity  # NEW
        if not rate_limiter.check_rate_limit(current_user_id, max_requests=100, window_seconds=3600):  # NEW
            return jsonify({  # NEW
                "error": "rate_limit_exceeded",  # NEW
                "message": "Too many requests. Please try again later."  # NEW
            }), 429  # NEW

        # Allow an optional JSON body, e.g. { "max_cases": 8 }  # NEW
        data = request.get_json(silent=True) or {}  # NEW
        max_cases = data.get("max_cases", 10)  # NEW

        service = EdgeCaseService()  # NEW
        edge_cases = service.generate_for_requirement(  # NEW
            requirement_id=requirement_id,  # NEW
            owner_id=current_user_id,       # NEW
            max_cases=max_cases,            # NEW
        )  # NEW

        return jsonify({  # NEW
            "requirement_id": requirement_id,  # NEW
            "edge_cases": edge_cases,          # NEW
        }), 201  # NEW

    except ValueError as e:  # NEW
        # e.g. requirement not found or access denied  # NEW
        return jsonify({"error": str(e)}), 404  # NEW
    except Exception as e:  # NEW
        print(f"Error generating edge cases: {str(e)}")  # NEW
        return jsonify({"error": "Failed to generate edge cases"}), 500  # NEW


# --- Ambiguity Detection Endpoints ---

@api_bp.route('/ambiguity/analyze', methods=['POST'])
@require_auth(["requirements:write"])
def analyze_ambiguity():
    """
    Analyze text for ambiguous terms.
    Expects JSON payload with 'text' and optional 'requirement_id'.
    """
    try:
        from flask import g
        current_user_id = g.user_id
        
        # Check rate limit
        if not rate_limiter.check_rate_limit(current_user_id, max_requests=1000, window_seconds=3600):
            return jsonify({
                "error": "rate_limit_exceeded",
                "message": "Too many requests. Please try again later."
            }), 429
        
        data = request.get_json()
        if not data:
            return jsonify({"error": "No data provided"}), 400
        
        # Validate request using Pydantic schema
        try:
            validated_data = AmbiguityAnalyzeRequest(**data)
        except ValidationError as e:
            return jsonify({
                "error": "validation_error",
                "message": "Invalid request data",
                "details": e.errors()
            }), 400
        
        # Initialize service and run analysis
        service = AmbiguityService()
        analysis = service.run_analysis(
            text=validated_data.text,
            requirement_id=validated_data.requirement_id,
            owner_id=current_user_id,
            use_llm=validated_data.use_llm
        )
        
        # Return analysis with terms
        return jsonify({
            "id": analysis.id,
            "requirement_id": analysis.requirement_id,
            "owner_id": analysis.owner_id,
            "original_text": analysis.original_text,
            "analyzed_at": analysis.analyzed_at.isoformat() if analysis.analyzed_at else None,
            "total_terms_flagged": analysis.total_terms_flagged,
            "terms_resolved": analysis.terms_resolved,
            "status": analysis.status,
            "terms": [
                {
                    "id": term.id,
                    "term": term.term,
                    "position_start": term.position_start,
                    "position_end": term.position_end,
                    "sentence_context": term.sentence_context,
                    "is_ambiguous": term.is_ambiguous,
                    "confidence": term.confidence,
                    "reasoning": term.reasoning,
                    "clarification_prompt": term.clarification_prompt,
                    "suggested_replacements": term.suggested_replacements,
                    "status": term.status
                }
                for term in analysis.terms
            ]
        }), 201
        
    except Exception as e:
        print(f"Error analyzing ambiguity: {str(e)}")
        return jsonify({"error": f"Failed to analyze text: {str(e)}"}), 500


@api_bp.route('/ambiguity/analysis/<int:analysis_id>', methods=['GET'])
@require_auth(["requirements:read"])
def get_ambiguity_analysis(analysis_id):
    """
    Retrieve specific analysis results by ID.
    Optimized with eager loading to prevent N+1 query problems.
    """
    try:
        from flask import g
        current_user_id = g.user_id
        
        # Use optimized query with eager loading
        from .database_optimization import get_analysis_with_terms
        analysis = get_analysis_with_terms(analysis_id, owner_id=current_user_id)
        
        if not analysis:
            return jsonify({"error": "Analysis not found or access denied"}), 404
        
        # Return analysis with terms
        return jsonify({
            "id": analysis.id,
            "requirement_id": analysis.requirement_id,
            "owner_id": analysis.owner_id,
            "original_text": analysis.original_text,
            "analyzed_at": analysis.analyzed_at.isoformat() if analysis.analyzed_at else None,
            "total_terms_flagged": analysis.total_terms_flagged,
            "terms_resolved": analysis.terms_resolved,
            "status": analysis.status,
            "terms": [
                {
                    "id": term.id,
                    "term": term.term,
                    "position_start": term.position_start,
                    "position_end": term.position_end,
                    "sentence_context": term.sentence_context,
                    "is_ambiguous": term.is_ambiguous,
                    "confidence": term.confidence,
                    "reasoning": term.reasoning,
                    "clarification_prompt": term.clarification_prompt,
                    "suggested_replacements": term.suggested_replacements,
                    "status": term.status
                }
                for term in analysis.terms
            ]
        })
        
    except Exception as e:
        print(f"Error retrieving analysis: {str(e)}")
        return jsonify({"error": f"Failed to retrieve analysis: {str(e)}"}), 500


@api_bp.route('/ambiguity/analyze/requirement/<int:requirement_id>', methods=['POST'])
@require_auth(["requirements:write"])
def analyze_requirement_ambiguity(requirement_id):
    """
    Analyze a specific requirement for ambiguous terms.
    """
    try:
        from flask import g
        current_user_id = g.user_id
        
        # Check rate limit
        if not rate_limiter.check_rate_limit(current_user_id, max_requests=1000, window_seconds=3600):
            return jsonify({
                "error": "rate_limit_exceeded",
                "message": "Too many requests. Please try again later."
            }), 429
        
        # Get JSON data, handle empty body gracefully
        try:
            data = request.get_json(silent=True) or {}
        except Exception:
            data = {}
        
        # Validate request using Pydantic schema
        try:
            validated_data = AmbiguityAnalyzeRequirementRequest(**data)
        except ValidationError as e:
            return jsonify({
                "error": "validation_error",
                "message": "Invalid request data",
                "details": e.errors()
            }), 400
        
        use_llm = validated_data.use_llm
        
        # Initialize service and run analysis
        service = AmbiguityService()
        analysis = service.run_requirement_analysis(
            requirement_id=requirement_id,
            owner_id=current_user_id,
            use_llm=use_llm
        )
        
        # Return analysis with terms
        return jsonify({
            "id": analysis.id,
            "requirement_id": analysis.requirement_id,
            "owner_id": analysis.owner_id,
            "original_text": analysis.original_text,
            "analyzed_at": analysis.analyzed_at.isoformat() if analysis.analyzed_at else None,
            "total_terms_flagged": analysis.total_terms_flagged,
            "terms_resolved": analysis.terms_resolved,
            "status": analysis.status,
            "terms": [
                {
                    "id": term.id,
                    "term": term.term,
                    "position_start": term.position_start,
                    "position_end": term.position_end,
                    "sentence_context": term.sentence_context,
                    "is_ambiguous": term.is_ambiguous,
                    "confidence": term.confidence,
                    "reasoning": term.reasoning,
                    "clarification_prompt": term.clarification_prompt,
                    "suggested_replacements": term.suggested_replacements,
                    "status": term.status
                }
                for term in analysis.terms
            ]
        }), 201
        
    except ValueError as e:
        return jsonify({"error": str(e)}), 404
    except Exception as e:
        print(f"Error analyzing requirement: {str(e)}")
        return jsonify({"error": f"Failed to analyze requirement: {str(e)}"}), 500


@api_bp.route('/ambiguity/analyze/batch', methods=['POST'])
@require_auth(["requirements:write"])
def analyze_batch_ambiguity():
    """
    Batch analyze multiple requirements for ambiguous terms.
    Expects JSON payload with 'requirement_ids' array.
    """
    try:
        from flask import g
        current_user_id = g.user_id
        
        # Check rate limit
        if not rate_limiter.check_rate_limit(current_user_id, max_requests=1000, window_seconds=3600):
            return jsonify({
                "error": "rate_limit_exceeded",
                "message": "Too many batch requests. Please try again later."
            }), 429
        
        data = request.get_json()
        if not data:
            return jsonify({"error": "No data provided"}), 400
        
        # Validate request using Pydantic schema
        try:
            validated_data = AmbiguityBatchAnalyzeRequest(**data)
        except ValidationError as e:
            return jsonify({
                "error": "validation_error",
                "message": "Invalid request data",
                "details": e.errors()
            }), 400
        
        requirement_ids = validated_data.requirement_ids
        use_llm = validated_data.use_llm
        
        # Initialize service and run batch analysis
        service = AmbiguityService()
        analyses = service.run_batch_analysis(
            requirement_ids=requirement_ids,
            owner_id=current_user_id,
            use_llm=use_llm
        )
        
        # Return list of analyses
        results = []
        for analysis in analyses:
            results.append({
                "id": analysis.id,
                "requirement_id": analysis.requirement_id,
                "owner_id": analysis.owner_id,
                "original_text": analysis.original_text,
                "analyzed_at": analysis.analyzed_at.isoformat() if analysis.analyzed_at else None,
                "total_terms_flagged": analysis.total_terms_flagged,
                "terms_resolved": analysis.terms_resolved,
                "status": analysis.status,
                "terms": [
                    {
                        "id": term.id,
                        "term": term.term,
                        "position_start": term.position_start,
                        "position_end": term.position_end,
                        "sentence_context": term.sentence_context,
                        "is_ambiguous": term.is_ambiguous,
                        "confidence": term.confidence,
                        "reasoning": term.reasoning,
                        "clarification_prompt": term.clarification_prompt,
                        "suggested_replacements": term.suggested_replacements,
                        "status": term.status
                    }
                    for term in analysis.terms
                ]
            })
        
        return jsonify({
            "total_analyzed": len(results),
            "analyses": results
        }), 201
        
    except Exception as e:
        print(f"Error in batch analysis: {str(e)}")
        return jsonify({"error": f"Failed to analyze batch: {str(e)}"}), 500



@api_bp.route('/ambiguity/clarify', methods=['POST'])
@require_auth(["requirements:write"])
def submit_clarification():
    """
    Submit clarification for an ambiguous term.
    Expects JSON payload with analysis_id, term_id, clarified_text, and action.
    """
    try:
        from flask import g
        current_user_id = g.user_id
        
        # Check rate limit
        if not rate_limiter.check_rate_limit(current_user_id, max_requests=1000, window_seconds=3600):
            return jsonify({
                "error": "rate_limit_exceeded",
                "message": "Too many requests. Please try again later."
            }), 429
        
        data = request.get_json()
        if not data:
            return jsonify({"error": "No data provided"}), 400
        
        # Validate request using Pydantic schema
        try:
            validated_data = ClarificationSubmitRequest(**data)
        except ValidationError as e:
            return jsonify({
                "error": "validation_error",
                "message": "Invalid request data",
                "details": e.errors()
            }), 400
        
        analysis_id = validated_data.analysis_id
        term_id = validated_data.term_id
        clarified_text = validated_data.clarified_text
        action = validated_data.action
        
        # Get analysis and verify ownership
        analysis = AmbiguityAnalysis.query.filter_by(
            id=analysis_id,
            owner_id=current_user_id
        ).first()
        
        if not analysis:
            return jsonify({"error": "Analysis not found or access denied"}), 404
        
        # Get term and verify it belongs to this analysis
        term = AmbiguousTerm.query.filter_by(
            id=term_id,
            analysis_id=analysis_id
        ).first()
        
        if not term:
            return jsonify({"error": "Term not found in this analysis"}), 404
        
        # Get requirement if associated
        requirement = None
        if analysis.requirement_id:
            requirement = Requirement.query.filter_by(
                id=analysis.requirement_id,
                owner_id=current_user_id
            ).first()
            
            if not requirement:
                return jsonify({"error": "Associated requirement not found or access denied"}), 404
        
        # Store original text
        original_text = analysis.original_text
        
        # Update requirement text if associated
        updated_requirement = None
        if requirement:
            if action == 'replace':
                # Replace the ambiguous term in the requirement
                # Update both title and description
                if term.term.lower() in requirement.title.lower():
                    requirement.title = requirement.title.replace(term.term, clarified_text)
                if requirement.description and term.term.lower() in requirement.description.lower():
                    requirement.description = requirement.description.replace(term.term, clarified_text)
            else:  # append
                # Append clarification to description
                clarification_note = f"\n\nClarification: '{term.term}' means {clarified_text}"
                if requirement.description:
                    requirement.description += clarification_note
                else:
                    requirement.description = clarification_note
            
            updated_requirement = {
                "id": requirement.id,
                "req_id": requirement.req_id,
                "title": requirement.title,
                "description": requirement.description,
                "status": requirement.status,
                "priority": requirement.priority,
                "stakeholders": requirement.stakeholders,
                "requirement_type": requirement.requirement_type
            }
        
        # Create clarification history record
        from .models import ClarificationHistory
        clarification = ClarificationHistory(
            term_id=term_id,
            requirement_id=analysis.requirement_id,
            owner_id=current_user_id,
            original_text=original_text,
            clarified_text=clarified_text,
            action=action,
            clarified_at=datetime.utcnow()
        )
        db.session.add(clarification)
        
        # Update term status
        term.status = 'clarified'
        
        # Update analysis progress
        analysis.terms_resolved = sum(1 for t in analysis.terms if t.status == 'clarified')
        if analysis.terms_resolved >= analysis.total_terms_flagged:
            analysis.status = 'completed'
        else:
            analysis.status = 'in_progress'
        
        db.session.commit()
        
        return jsonify({
            "message": "Clarification submitted successfully",
            "clarification": {
                "id": clarification.id,
                "term_id": term_id,
                "clarified_text": clarified_text,
                "action": action,
                "clarified_at": clarification.clarified_at.isoformat()
            },
            "analysis": {
                "id": analysis.id,
                "terms_resolved": analysis.terms_resolved,
                "total_terms_flagged": analysis.total_terms_flagged,
                "status": analysis.status
            },
            "updated_requirement": updated_requirement
        })
        
    except Exception as e:
        db.session.rollback()
        print(f"Error submitting clarification: {str(e)}")
        return jsonify({"error": f"Failed to submit clarification: {str(e)}"}), 500


@api_bp.route('/ambiguity/suggestions/<int:term_id>', methods=['GET'])
@require_auth(["requirements:read"])
def get_term_suggestions(term_id):
    """
    Get AI-generated suggestions for a specific ambiguous term.
    """
    try:
        from flask import g
        current_user_id = g.user_id
        
        # Get term
        term = AmbiguousTerm.query.filter_by(id=term_id).first()
        
        if not term:
            return jsonify({"error": "Term not found"}), 404
        
        # Verify ownership through analysis
        analysis = AmbiguityAnalysis.query.filter_by(
            id=term.analysis_id,
            owner_id=current_user_id
        ).first()
        
        if not analysis:
            return jsonify({"error": "Access denied"}), 403
        
        # Return suggestions and prompt
        return jsonify({
            "term_id": term.id,
            "term": term.term,
            "sentence_context": term.sentence_context,
            "suggestions": term.suggested_replacements or [],
            "clarification_prompt": term.clarification_prompt or f"What specific criteria do you mean by '{term.term}'?",
            "confidence": term.confidence,
            "reasoning": term.reasoning
        })
        
    except Exception as e:
        print(f"Error retrieving suggestions: {str(e)}")
        return jsonify({"error": f"Failed to retrieve suggestions: {str(e)}"}), 500



@api_bp.route('/ambiguity/report/<int:requirement_id>', methods=['GET'])
@require_auth(["requirements:read"])
def get_requirement_ambiguity_report(requirement_id):
    """
    Generate ambiguity report for a specific requirement.
    """
    try:
        from flask import g
        current_user_id = g.user_id
        
        # Verify requirement ownership
        requirement = Requirement.query.filter_by(
            id=requirement_id,
            owner_id=current_user_id
        ).first()
        
        if not requirement:
            return jsonify({"error": "Requirement not found or access denied"}), 404
        
        # Get all analyses for this requirement
        analyses = AmbiguityAnalysis.query.filter_by(
            requirement_id=requirement_id,
            owner_id=current_user_id
        ).order_by(AmbiguityAnalysis.analyzed_at.desc()).all()
        
        if not analyses:
            return jsonify({
                "requirement_id": requirement_id,
                "requirement_title": requirement.title,
                "total_analyses": 0,
                "total_terms_flagged": 0,
                "total_terms_resolved": 0,
                "status": "no_analysis",
                "analyses": []
            })
        
        # Get latest analysis
        latest_analysis = analyses[0]
        
        # Aggregate statistics
        total_terms_flagged = sum(a.total_terms_flagged for a in analyses)
        total_terms_resolved = sum(a.terms_resolved for a in analyses)
        
        # Build report
        report = {
            "requirement_id": requirement_id,
            "requirement_title": requirement.title,
            "requirement_description": requirement.description,
            "total_analyses": len(analyses),
            "total_terms_flagged": latest_analysis.total_terms_flagged,
            "total_terms_resolved": latest_analysis.terms_resolved,
            "resolution_percentage": round((latest_analysis.terms_resolved / latest_analysis.total_terms_flagged * 100) if latest_analysis.total_terms_flagged > 0 else 0, 1),
            "status": latest_analysis.status,
            "latest_analysis_date": latest_analysis.analyzed_at.isoformat() if latest_analysis.analyzed_at else None,
            "terms": [
                {
                    "id": term.id,
                    "term": term.term,
                    "sentence_context": term.sentence_context,
                    "status": term.status,
                    "confidence": term.confidence,
                    "reasoning": term.reasoning,
                    "clarification_prompt": term.clarification_prompt,
                    "suggested_replacements": term.suggested_replacements,
                    "clarifications": [
                        {
                            "clarified_text": c.clarified_text,
                            "action": c.action,
                            "clarified_at": c.clarified_at.isoformat()
                        }
                        for c in term.clarifications
                    ]
                }
                for term in latest_analysis.terms
            ]
        }
        
        return jsonify(report)
        
    except Exception as e:
        print(f"Error generating requirement report: {str(e)}")
        return jsonify({"error": f"Failed to generate report: {str(e)}"}), 500


@api_bp.route('/ambiguity/report/project', methods=['GET'])
@require_auth(["requirements:read"])
def get_project_ambiguity_report():
    """
    Generate project-wide ambiguity report for all user requirements.
    """
    try:
        from flask import g
        current_user_id = g.user_id
        
        # Get all analyses for this user
        analyses = AmbiguityAnalysis.query.filter_by(
            owner_id=current_user_id
        ).order_by(AmbiguityAnalysis.analyzed_at.desc()).all()
        
        if not analyses:
            return jsonify({
                "total_requirements_analyzed": 0,
                "total_terms_flagged": 0,
                "total_terms_resolved": 0,
                "resolution_percentage": 0,
                "requirements": []
            })
        
        # Group analyses by requirement
        requirement_analyses = {}
        for analysis in analyses:
            req_id = analysis.requirement_id
            if req_id not in requirement_analyses:
                requirement_analyses[req_id] = []
            requirement_analyses[req_id].append(analysis)
        
        # Build requirement summaries
        requirement_summaries = []
        total_terms_flagged = 0
        total_terms_resolved = 0
        
        for req_id, req_analyses in requirement_analyses.items():
            latest = req_analyses[0]  # Most recent analysis
            
            requirement = Requirement.query.filter_by(id=req_id).first()
            
            total_terms_flagged += latest.total_terms_flagged
            total_terms_resolved += latest.terms_resolved
            
            requirement_summaries.append({
                "requirement_id": req_id,
                "requirement_title": requirement.title if requirement else "Unknown",
                "req_id": requirement.req_id if requirement else None,
                "total_terms_flagged": latest.total_terms_flagged,
                "terms_resolved": latest.terms_resolved,
                "resolution_percentage": round((latest.terms_resolved / latest.total_terms_flagged * 100) if latest.total_terms_flagged > 0 else 0, 1),
                "status": latest.status,
                "last_analyzed": latest.analyzed_at.isoformat() if latest.analyzed_at else None,
                "pending_terms": [
                    {
                        "term": term.term,
                        "sentence_context": term.sentence_context,
                        "confidence": term.confidence
                    }
                    for term in latest.terms if term.status == 'pending'
                ]
            })
        
        # Sort by most terms flagged
        requirement_summaries.sort(key=lambda x: x['total_terms_flagged'], reverse=True)
        
        # Build project report
        report = {
            "total_requirements_analyzed": len(requirement_analyses),
            "total_terms_flagged": total_terms_flagged,
            "total_terms_resolved": total_terms_resolved,
            "resolution_percentage": round((total_terms_resolved / total_terms_flagged * 100) if total_terms_flagged > 0 else 0, 1),
            "requirements": requirement_summaries
        }
        
        return jsonify(report)
        
    except Exception as e:
        print(f"Error generating project report: {str(e)}")
        return jsonify({"error": f"Failed to generate project report: {str(e)}"}), 500


@api_bp.route('/ambiguity/report/export', methods=['POST'])
@require_auth(["requirements:read"])
def export_ambiguity_report():
    """
    Export ambiguity report in specified format (.txt or .md).
    Expects JSON payload with 'requirement_ids' array and 'format'.
    """
    try:
        from flask import g, make_response
        current_user_id = g.user_id
        
        # Check rate limit
        if not rate_limiter.check_rate_limit(current_user_id, max_requests=50, window_seconds=3600):
            return jsonify({
                "error": "rate_limit_exceeded",
                "message": "Too many export requests. Please try again later."
            }), 429
        
        data = request.get_json()
        if not data:
            return jsonify({"error": "No data provided"}), 400
        
        # Validate request using Pydantic schema
        try:
            validated_data = ReportExportRequest(**data)
        except ValidationError as e:
            return jsonify({
                "error": "validation_error",
                "message": "Invalid request data",
                "details": e.errors()
            }), 400
        
        requirement_ids = validated_data.requirement_ids or []
        export_format = validated_data.format
        
        # If no specific requirements, export all
        if not requirement_ids:
            analyses = AmbiguityAnalysis.query.filter_by(
                owner_id=current_user_id
            ).order_by(AmbiguityAnalysis.analyzed_at.desc()).all()
            
            # Get unique requirement IDs
            requirement_ids = list(set(a.requirement_id for a in analyses if a.requirement_id))
        
        # Build report content
        report_lines = []
        
        if export_format == 'md':
            report_lines.append("# Ambiguity Detection Report\n")
            report_lines.append(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}\n\n")
        else:
            report_lines.append("AMBIGUITY DETECTION REPORT\n")
            report_lines.append("=" * 50 + "\n")
            report_lines.append(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S UTC')}\n\n")
        
        # Process each requirement
        for req_id in requirement_ids:
            requirement = Requirement.query.filter_by(
                id=req_id,
                owner_id=current_user_id
            ).first()
            
            if not requirement:
                continue
            
            # Get latest analysis
            analysis = AmbiguityAnalysis.query.filter_by(
                requirement_id=req_id,
                owner_id=current_user_id
            ).order_by(AmbiguityAnalysis.analyzed_at.desc()).first()
            
            if not analysis:
                continue
            
            # Format requirement section
            if export_format == 'md':
                report_lines.append(f"## {requirement.req_id}: {requirement.title}\n\n")
                report_lines.append(f"**Status:** {analysis.status}\n")
                report_lines.append(f"**Terms Flagged:** {analysis.total_terms_flagged}\n")
                report_lines.append(f"**Terms Resolved:** {analysis.terms_resolved}\n")
                report_lines.append(f"**Resolution:** {round((analysis.terms_resolved / analysis.total_terms_flagged * 100) if analysis.total_terms_flagged > 0 else 0, 1)}%\n\n")
                
                if analysis.terms:
                    report_lines.append("### Ambiguous Terms\n\n")
                    for term in analysis.terms:
                        report_lines.append(f"#### {term.term} ({term.status})\n\n")
                        report_lines.append(f"**Context:** {term.sentence_context}\n\n")
                        report_lines.append(f"**Confidence:** {term.confidence}\n\n")
                        if term.reasoning:
                            report_lines.append(f"**Reasoning:** {term.reasoning}\n\n")
                        if term.suggested_replacements:
                            report_lines.append("**Suggestions:**\n")
                            for suggestion in term.suggested_replacements:
                                report_lines.append(f"- {suggestion}\n")
                            report_lines.append("\n")
                        if term.clarifications:
                            report_lines.append("**Clarifications:**\n")
                            for clarification in term.clarifications:
                                report_lines.append(f"- {clarification.clarified_text} ({clarification.action}) - {clarification.clarified_at.strftime('%Y-%m-%d')}\n")
                            report_lines.append("\n")
                report_lines.append("---\n\n")
            else:
                report_lines.append(f"\n{requirement.req_id}: {requirement.title}\n")
                report_lines.append("-" * 50 + "\n")
                report_lines.append(f"Status: {analysis.status}\n")
                report_lines.append(f"Terms Flagged: {analysis.total_terms_flagged}\n")
                report_lines.append(f"Terms Resolved: {analysis.terms_resolved}\n")
                report_lines.append(f"Resolution: {round((analysis.terms_resolved / analysis.total_terms_flagged * 100) if analysis.total_terms_flagged > 0 else 0, 1)}%\n\n")
                
                if analysis.terms:
                    report_lines.append("Ambiguous Terms:\n\n")
                    for term in analysis.terms:
                        report_lines.append(f"  {term.term} ({term.status})\n")
                        report_lines.append(f"  Context: {term.sentence_context}\n")
                        report_lines.append(f"  Confidence: {term.confidence}\n")
                        if term.reasoning:
                            report_lines.append(f"  Reasoning: {term.reasoning}\n")
                        if term.suggested_replacements:
                            report_lines.append("  Suggestions:\n")
                            for suggestion in term.suggested_replacements:
                                report_lines.append(f"    - {suggestion}\n")
                        if term.clarifications:
                            report_lines.append("  Clarifications:\n")
                            for clarification in term.clarifications:
                                report_lines.append(f"    - {clarification.clarified_text} ({clarification.action})\n")
                        report_lines.append("\n")
        
        # Create response
        report_content = "".join(report_lines)
        response = make_response(report_content)
        response.headers['Content-Type'] = 'text/plain' if export_format == 'txt' else 'text/markdown'
        response.headers['Content-Disposition'] = f'attachment; filename=ambiguity_report.{export_format}'
        
        return response
        
    except Exception as e:
        print(f"Error exporting report: {str(e)}")
        return jsonify({"error": f"Failed to export report: {str(e)}"}), 500



@api_bp.route('/ambiguity/lexicon', methods=['GET'])
@require_auth(["requirements:read"])
def get_ambiguity_lexicon():
    """
    Get current lexicon (global + user-specific).
    Returns both global terms and user's custom terms.
    """
    try:
        from flask import g
        current_user_id = g.user_id
        
        # Initialize lexicon manager
        from .lexicon_manager import LexiconManager
        lexicon_manager = LexiconManager()
        
        # Get combined lexicon
        combined_lexicon = lexicon_manager.get_lexicon(owner_id=current_user_id)
        
        # Get user's custom terms separately
        custom_terms = lexicon_manager.get_user_custom_terms(owner_id=current_user_id)
        
        # Get global terms
        global_terms = lexicon_manager.get_default_lexicon()
        
        return jsonify({
            "global": sorted(global_terms),
            "custom_include": sorted(custom_terms['include']),
            "custom_exclude": sorted(custom_terms['exclude']),
            "combined": sorted(combined_lexicon),
            "total_terms": len(combined_lexicon)
        })
        
    except Exception as e:
        print(f"Error retrieving lexicon: {str(e)}")
        return jsonify({"error": f"Failed to retrieve lexicon: {str(e)}"}), 500


@api_bp.route('/ambiguity/lexicon/add', methods=['POST'])
@require_auth(["requirements:write"])
def add_lexicon_term():
    """
    Add term to user's custom lexicon.
    Expects JSON payload with 'term' and 'type' ('include' or 'exclude').
    """
    try:
        from flask import g
        current_user_id = g.user_id
        
        # Check rate limit
        if not rate_limiter.check_rate_limit(current_user_id, max_requests=1000, window_seconds=3600):
            return jsonify({
                "error": "rate_limit_exceeded",
                "message": "Too many requests. Please try again later."
            }), 429
        
        data = request.get_json()
        if not data:
            return jsonify({"error": "No data provided"}), 400
        
        # Validate request using Pydantic schema
        try:
            validated_data = LexiconAddRequest(**data)
        except ValidationError as e:
            return jsonify({
                "error": "validation_error",
                "message": "Invalid request data",
                "details": e.errors()
            }), 400
        
        term = validated_data.term
        term_type = validated_data.type
        category = validated_data.category
        
        # Map to database type
        db_type = f'custom_{term_type}'
        
        # Initialize lexicon manager
        from .lexicon_manager import LexiconManager
        lexicon_manager = LexiconManager()
        
        # Add term
        success = lexicon_manager.add_term(
            term=term,
            owner_id=current_user_id,
            term_type=db_type,
            category=category
        )
        
        if not success:
            return jsonify({"error": "Term already exists in lexicon"}), 409
        
        # Get updated lexicon
        updated_lexicon = lexicon_manager.get_lexicon(owner_id=current_user_id)
        custom_terms = lexicon_manager.get_user_custom_terms(owner_id=current_user_id)
        
        return jsonify({
            "message": f"Term '{term}' added to {term_type} list",
            "term": term,
            "type": term_type,
            "custom_include": sorted(custom_terms['include']),
            "custom_exclude": sorted(custom_terms['exclude']),
            "combined": sorted(updated_lexicon)
        }), 201
        
    except Exception as e:
        print(f"Error adding lexicon term: {str(e)}")
        return jsonify({"error": f"Failed to add term: {str(e)}"}), 500


@api_bp.route('/ambiguity/lexicon/<term>', methods=['DELETE'])
@require_auth(["requirements:write"])
def remove_lexicon_term(term):
    """
    Remove term from user's custom lexicon.
    Query parameter 'type' specifies 'include' or 'exclude' (default: 'include').
    """
    try:
        from flask import g
        current_user_id = g.user_id
        
        term_type = request.args.get('type', 'include')
        
        if term_type not in ['include', 'exclude']:
            return jsonify({"error": "type must be 'include' or 'exclude'"}), 400
        
        # Map to database type
        db_type = f'custom_{term_type}'
        
        # Initialize lexicon manager
        from .lexicon_manager import LexiconManager
        lexicon_manager = LexiconManager()
        
        # Remove term
        success = lexicon_manager.remove_term(
            term=term,
            owner_id=current_user_id,
            term_type=db_type
        )
        
        if not success:
            return jsonify({"error": "Term not found in lexicon"}), 404
        
        # Get updated lexicon
        updated_lexicon = lexicon_manager.get_lexicon(owner_id=current_user_id)
        custom_terms = lexicon_manager.get_user_custom_terms(owner_id=current_user_id)
        
        return jsonify({
            "message": f"Term '{term}' removed from {term_type} list",
            "term": term,
            "type": term_type,
            "custom_include": sorted(custom_terms['include']),
            "custom_exclude": sorted(custom_terms['exclude']),
            "combined": sorted(updated_lexicon)
        })
        
    except Exception as e:
        print(f"Error removing lexicon term: {str(e)}")
        return jsonify({"error": f"Failed to remove term: {str(e)}"}), 500


# --- Database Performance Monitoring Endpoints ---

@api_bp.route('/admin/database/stats', methods=['GET'])
@require_auth(["admin:read"])
def get_database_stats():
    """
    Get database performance statistics including table sizes and row counts.
    Admin only endpoint.
    """
    try:
        from .database_optimization import get_table_statistics
        
        stats = get_table_statistics()
        
        if not stats.get('success'):
            return jsonify({"error": stats.get('error', 'Failed to retrieve statistics')}), 500
        
        return jsonify(stats)
        
    except Exception as e:
        print(f"Error retrieving database stats: {str(e)}")
        return jsonify({"error": f"Failed to retrieve database stats: {str(e)}"}), 500


@api_bp.route('/admin/database/indexes', methods=['GET'])
@require_auth(["admin:read"])
def get_index_usage():
    """
    Get index usage statistics to identify unused or inefficient indexes.
    Admin only endpoint.
    """
    try:
        from .database_optimization import get_index_usage_statistics
        
        stats = get_index_usage_statistics()
        
        if not stats.get('success'):
            return jsonify({"error": stats.get('error', 'Failed to retrieve index statistics')}), 500
        
        return jsonify(stats)
        
    except Exception as e:
        print(f"Error retrieving index stats: {str(e)}")
        return jsonify({"error": f"Failed to retrieve index stats: {str(e)}"}), 500


@api_bp.route('/admin/database/pool', methods=['GET'])
@require_auth(["admin:read"])
def get_connection_pool_stats():
    """
    Get current connection pool statistics.
    Admin only endpoint.
    """
    try:
        from .database_optimization import get_connection_pool_stats
        
        stats = get_connection_pool_stats()
        
        if not stats.get('success'):
            return jsonify({"error": stats.get('error', 'Failed to retrieve pool statistics')}), 500
        
        return jsonify(stats)
        
    except Exception as e:
        print(f"Error retrieving pool stats: {str(e)}")
        return jsonify({"error": f"Failed to retrieve pool stats: {str(e)}"}), 500


@api_bp.route('/admin/database/query-stats', methods=['GET'])
@require_auth(["admin:read"])
def get_query_performance_stats():
    """
    Get query performance statistics including slow query information.
    Admin only endpoint.
    """
    try:
        from .database_optimization import query_monitor
        
        stats = query_monitor.get_stats()
        
        return jsonify(stats)
        
    except Exception as e:
        print(f"Error retrieving query stats: {str(e)}")
        return jsonify({"error": f"Failed to retrieve query stats: {str(e)}"}), 500


@api_bp.route('/admin/database/query-stats', methods=['DELETE'])
@require_auth(["admin:write"])
def clear_query_performance_stats():
    """
    Clear collected query performance statistics.
    Admin only endpoint.
    """
    try:
        from .database_optimization import query_monitor
        
        query_monitor.clear_stats()
        
        return jsonify({"message": "Query statistics cleared successfully"})
        
    except Exception as e:
        print(f"Error clearing query stats: {str(e)}")
        return jsonify({"error": f"Failed to clear query stats: {str(e)}"}), 500

@api_bp.route('/admin/database/analyze-query', methods=['POST'])
@require_auth(["admin:write"])
def analyze_query():
    """
    Run EXPLAIN ANALYZE on a provided SQL query.
    Admin only endpoint for debugging query performance.
    Expects JSON payload with 'query' and optional 'params'.
    """
    try:
        data = request.get_json()
        if not data or 'query' not in data:
            return jsonify({"error": "No query provided"}), 400
        
        query_sql = data['query']
        params = data.get('params', {})
        
        from .database_optimization import analyze_query_plan
        
        result = analyze_query_plan(query_sql, params)
        
        if not result.get('success'):
            return jsonify({"error": result.get('error', 'Query analysis failed')}), 500
        
        return jsonify(result)
        
    except Exception as e:
        print(f"Error analyzing query: {str(e)}")
        return jsonify({"error": f"Failed to analyze query: {str(e)}"}), 500

@api_bp.route('/documents/<int:document_id>/analyze/contradictions', methods=['POST'])
@require_auth(["documents:write"])
def trigger_contradiction_analysis(document_id):
    """
    Triggers the LLM-based contradiction analysis for all requirements 
    linked to a specific document ID.
    """
    from flask import g
    current_user_id = g.user_id 

    document = Document.query.filter_by(id=document_id, owner_id=current_user_id).first()
    if not document:
        return jsonify({"message": "Document not found or access denied."}), 404
        
    try:
        data = request.get_json(silent=True) or {}
        project_context = data.get('project_context')
        

        analysis_service = ContradictionAnalysisService(db_instance=db, user_id=current_user_id)
        new_report = analysis_service.run_analysis(
            document_id=document_id, 
            project_context=project_context
        )
        
        # Use the schema to serialize the new report for the API response
        schema = ContradictionAnalysisSchema()
        return jsonify(schema.model_dump(new_report)), 200

    except ValueError as e:
        # Handle cases where no requirements are found
        return jsonify({"error": str(e)}), 400
    except Exception as e:
        # Generic error handling
        db.session.rollback()
        print(f"Error during contradiction analysis: {e}")
        return jsonify({"error": "Failed to complete contradiction analysis."}), 500

@api_bp.route('/documents/<int:document_id>/analyze/contradictions/latest', methods=['GET'])
@require_auth(["documents:read"])
def get_latest_contradiction_report(document_id):
    """
    Retrieves the most recently run contradiction analysis report for a document.
    """
    from flask import g
    current_user_id = g.user_id
    document = Document.query.filter_by(id=document_id, owner_id=current_user_id).first()
    if not document:
        return jsonify({"message": "Document not found or access denied."}), 404
    # [ --- END FIX --- ]
    
    analysis_service = ContradictionAnalysisService(db_instance=db)
    report = analysis_service.get_latest_analysis(document_id)
    
    if not report:
        # This is a 200 OK, not a 404, because the *document* exists, just no report.
        return jsonify({"message": "No contradiction analysis report found for this document."}), 200
        
    schema = ContradictionAnalysisSchema()
    return jsonify(schema.model_dump(report)), 200

# [ --- ADD THIS NEW ROUTE --- ]

@api_bp.route('/project/analyze/contradictions', methods=['POST'])
@require_auth(["documents:write"]) # Use same permission as single analysis
def project_contradiction_analysis():
    """
    Runs contradiction analysis on ALL documents for the current user.
    Aggregates all conflicts into a single report.
    """
    from flask import g
    current_user_id = g.user_id
    
    try:
        # 1. Fetch all documents for that user
        documents = Document.query.filter_by(owner_id=current_user_id).all()
        if not documents:
            # No documents, return an empty 'complete' report
            return jsonify({
                "status": "complete", 
                "conflicts": [], 
                "total_conflicts_found": 0,
                "analyzed_at": datetime.utcnow().isoformat()
            }), 200

        # 2. Create an empty list to hold all conflict objects
        all_conflicts = []
        
        # 3. Instantiate the service
        analysis_service = ContradictionAnalysisService(db_instance=db, user_id=current_user_id)
        
        # 4. Loop through each doc and run the internal analysis
        for doc in documents:
            print(f"Running project analysis on: Document {doc.id} ({doc.filename})")
            try:
                # 5. Run your internal analysis logic
                # This creates a new ContradictionAnalysis and its ConflictingPairs in the DB
                new_report = analysis_service.run_analysis(
                    document_id=doc.id, 
                    project_context=None # Context is per-document, so None for a batch job
                )
                
                # 6. Append conflicts found to the main list
                all_conflicts.extend(new_report.conflicts)
                
            except ValueError as ve:
                # This is raised if a document has no requirements
                print(f"Skipping document {doc.id}: {str(ve)}")
            except Exception as e:
                # Don't let one failed document stop the whole project
                print(f"Error analyzing document {doc.id}: {str(e)}")
        
        # 7. Convert ConflictingPair objects to dictionaries
        conflict_dicts = [
            {
                "id": c.id,
                "analysis_id": c.analysis_id,
                "conflict_id": c.conflict_id,
                "reason": c.reason,
                "conflicting_requirement_ids": c.conflicting_requirement_ids,
                "status": c.status
            } for c in all_conflicts
        ]
        
        # 8. Finally, return the aggregated report
        return jsonify({
            "status": "complete",
            "conflicts": conflict_dicts,
            "total_conflicts_found": len(conflict_dicts),
            "analyzed_at": datetime.utcnow().isoformat()
        }), 200

    except Exception as e:
        db.session.rollback() # Rollback any partial commits if the outer loop fails
        print(f"Error in project-wide contradiction analysis: {e}")
        return jsonify({"error": "Failed to run project analysis", "message": str(e)}), 500

# [ --- END OF NEW ROUTE --- ]


@api_bp.route('/documents/<int:document_id>/requirements', methods=['GET'])
@require_auth(["requirements:read"])
def get_document_requirements(document_id):
    """
    Fetches all requirements associated with a specific document.
    """
    try:
        from flask import g
        current_user_id = g.user_id
        
        # Check if document exists and is owned by user
        document = Document.query.filter_by(id=document_id, owner_id=current_user_id).first()
        if not document:
            return jsonify({"error": "Document not found or access denied"}), 404

        # Get requirements linked to this document
        requirements = Requirement.query.filter_by(
            source_document_id=document_id,
            owner_id=current_user_id
        ).all()

        results = [
            {
                "id": req.id,
                "req_id": req.req_id,
                "title": req.title,
                "description": req.description,
                "status": req.status,
                "priority": req.priority,
                "stakeholders": req.stakeholders,
                "requirement_type": req.requirment_type,
                "source_document_filename": document.filename, # We already have the doc
                "tags": [{"id": tag.id, "name": tag.name} for tag in req.tags]
            }
            for req in requirements
        ]
        
        return jsonify({"requirements": results}) 
    
    except Exception as e:
        print(f"Error fetching document requirements: {str(e)}")
        return jsonify({"error": "Failed to fetch document requirements"}), 500
    
@api_bp.route('/requirements/<int:requirement_id>', methods=['PUT'])
@require_auth(["requirements:write"])
def update_requirement(requirement_id):
    """
    Updates a single requirement's details (e.g., title, description, status).
    """
    try:
        from flask import g
        current_user_id = g.user_id
        
        # Security: Find the requirement AND verify ownership
        requirement = Requirement.query.filter_by(
            id=requirement_id, 
            owner_id=current_user_id
        ).first()

        if not requirement:
            return jsonify({"error": "Requirement not found or access denied"}), 404

        data = request.get_json()
        if not data:
            return jsonify({"error": "No update data provided"}), 400

        # Update fields only if they are provided in the request
        # This is a PATCH-like behavior, which is good for frontend flexibility
        requirement.title = data.get('title', requirement.title)
        requirement.description = data.get('description', requirement.description)
        requirement.status = data.get('status', requirement.status)
        requirement.priority = data.get('priority', requirement.priority)
        requirement.stakeholders = data.get('stakeholders', requirement.stakeholders)
        requirement.requirement_type = data.get('requirement_type', requirement.requirement_type)
        
        # Note: Updating tags is more complex (many-to-many)
        # We'll skip it for this basic update.

        db.session.commit()

        # Best practice: return the updated JSON of the requirement
        updated_data = {
            "id": requirement.id,
            "req_id": requirement.req_id,
            "title": requirement.title,
            "description": requirement.description,
            "status": requirement.status,
            "priority": requirement.priority,
            "stakeholders": requirement.stakeholders,
            "requirement_type": requirement.requirement_type,
            "source_document_filename": requirement.source_document.filename if requirement.source_document else None,
            "tags": [{"id": tag.id, "name": tag.name} for tag in requirement.tags]
        }
        return jsonify(updated_data), 200

    except Exception as e:
        db.session.rollback()
        print(f"Error updating requirement: {str(e)}")
        return jsonify({"error": f"Failed to update requirement: {str(e)}"}), 500


@api_bp.route('/requirements/<int:requirement_id>', methods=['DELETE'])
@require_auth(["requirements:write"])
def delete_requirement(requirement_id):
    """
    Deletes a single requirement.
    """
    try:
        from flask import g
        current_user_id = g.user_id
        
        # Security: Find the requirement AND verify ownership
        requirement = Requirement.query.filter_by(
            id=requirement_id, 
            owner_id=current_user_id
        ).first()

        if not requirement:
            return jsonify({"error": "Requirement not found or access denied"}), 404

        # Delete the requirement
        db.session.delete(requirement)
        db.session.commit()

        return jsonify({"message": f"Requirement {requirement.req_id} deleted successfully"}), 200

    except Exception as e:
        db.session.rollback()
        print(f"Error deleting requirement: {str(e)}")
        return jsonify({"error": f"Failed to delete requirement: {str(e)}"}), 500
  